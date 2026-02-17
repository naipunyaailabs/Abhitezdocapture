from app.services.llm_service import llm_service
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import io

class RFPService:
    RFP_CREATION_PROMPT = """You are an elite Proposal Manager & Bid Writer.
Task: Create a winning, compliant, and highly professional Request for Proposal (RFP) document.

OBJECTIVES:
1. **Professionalism**: Use formal, industry-standard business language.
2. **Completeness**: Elaborate on every section. Avoid brief or generic descriptions.
3. **Structure**: Use clear hierarchy (I, 1.1, a).
4. **Clarity**: Define requirements precisely to minimize vendor questions.

CONTENT RULES:
- If specific details are missing, use standard industry best practices to fill the gaps or use square brackets [Insert specific...] ONLY if absolutely necessary.
- **Do not** use generic placeholders like "Lorem ipsum". Write actual, usable content.
- Ensure the 'Evaluation Criteria' section is robust and quantifiable.

RFP DETAILS:
Title: {{title}}
Organization: {{organization}}
Deadline: {{deadline}}

REQUIRED SECTIONS:
{{sections}}

OUTPUT FORMAT:
Strict Markdown.
"""

    async def create_rfp(self, title: str, organization: str, deadline: str, sections: list) -> dict:
        sections_text = "\n".join([f"- {s.get('title', '')}: {s.get('content', 'Draft valid professional content for this section')}" for s in sections])
        
        prompt = self.RFP_CREATION_PROMPT.replace('{{title}}', title) \
            .replace('{{organization}}', organization) \
            .replace('{{deadline}}', deadline) \
            .replace('{{sections}}', sections_text)

        system_prompt = "You are an expert Proposal Writer. Generate a comprehensive, professional RFP document in Markdown format. Focus on clarity, compliance, and detail."

        print(f"[RFPService] Creating RFP: {title}")
        response = await llm_service.unified_chat_completion(system_prompt, prompt)
        
        # Clean markdown code blocks
        content = response.strip()
        if content.startswith("```"):
            content = content.replace("```markdown", "").replace("```", "").strip()

        enhanced_sections = self.parse_markdown_response(content, sections)

        return {
            "title": title,
            "organization": organization,
            "deadline": deadline,
            "sections": enhanced_sections
        }

    async def create_standard_rfp(self, title: str, organization: str, deadline: str) -> dict:
        standard_sections = [
            {"title": "Executive Summary", "content": "Provide a comprehensive overview of the project..."},
            {"title": "Project Background and Objectives", "content": "Detail the background information..."},
            {"title": "Scope of Work", "content": "Define the detailed scope of work..."},
            {"title": "Technical Requirements", "content": "List all technical specifications..."},
            {"title": "Submission Requirements", "content": "Specify detailed requirements for proposal submissions..."},
            {"title": "Evaluation Criteria and Scoring", "content": "Detail the comprehensive evaluation process..."},
            {"title": "Project Timeline and Milestones", "content": "Provide a detailed project timeline..."},
            {"title": "Terms and Conditions", "content": "Include all contractual terms..."},
            {"title": "Budget and Pricing Structure", "content": "Specify the available budget range..."},
            {"title": "Vendor Qualifications and Experience", "content": "Define the required qualifications..."}
        ]
        
        return await self.create_rfp(title, organization, deadline, standard_sections)

    def parse_markdown_response(self, markdown_content: str, original_sections: list) -> list:
        if not markdown_content:
            return original_sections

        # Simple parser to map sections back
        # This is a basic implementation; a more robust one would use regex or a markdown parser ast
        lines = markdown_content.split('\n')
        section_map = {}
        current_title = None
        current_content = []

        for line in lines:
            line = line.strip()
            # Detect headers ## or ###
            # Original TS code regex: /^#{2,3}\s+(.+)$/
            match = re.match(r'^#{2,3}\s+(.+)$', line)
            if match:
                if current_title:
                    section_map[current_title] = "\n".join(current_content).strip()
                current_title = match.group(1).strip()
                current_content = []
            elif current_title:
                current_content.append(line)
        
        if current_title:
            section_map[current_title] = "\n".join(current_content).strip()

        enhanced_sections = []
        for section in original_sections:
            title = section['title']
            content = section_map.get(title)
            # Fuzzy match or robust mapping could be added here
            # For now, exact match or fallback
            if content and len(content) > len(section.get('content', '')):
                 enhanced_sections.append({"title": title, "content": content})
            else:
                enhanced_sections.append(section)
        
        # If we missed sections from the LLM response that weren't in original, we might want to add them
        # But for now, sticking to the original list structure as per TS code logic mostly.
        
        return enhanced_sections

    def create_rfp_word_document(self, rfp_content: dict) -> bytes:
        doc = Document()
        
        # Title
        title = doc.add_heading(rfp_content.get('title', 'Request for Proposal'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Prepared by: {rfp_content.get('organization', '')}")
        doc.add_paragraph(f"Deadline: {rfp_content.get('deadline', '')}")
        doc.add_page_break()
        
        # Table of Contents (Manual placeholder as Word fields are complex in python-docx)
        doc.add_heading("Table of Contents", level=1)
        for i, section in enumerate(rfp_content.get('sections', [])):
            doc.add_paragraph(f"{i+1}. {section['title']}")
        doc.add_page_break()

        # Sections
        for section in rfp_content.get('sections', []):
            doc.add_heading(section['title'], level=1)
            # Add content paragraphs
            # Basic markdown handling: 
            # - Remove ** for bold (or handle it)
            # - Handle bullet points
            lines = section['content'].split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Bullet points
                if line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                # Numbered lists (simple heuristic)
                elif re.match(r'^\d+\.\s', line):
                     p = doc.add_paragraph(line.split(' ', 1)[1], style='List Number')
                # Headers in content (###)
                elif line.startswith('###'):
                    doc.add_heading(line.replace('#', '').strip(), level=2)
                # Normal paragraph
                else:
                    p = doc.add_paragraph()
                    # Bold text handling **text**
                    # Split by ** and alternate bold/normal
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1: # Odd parts are between **
                            run.bold = True
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

rfp_service = RFPService()
