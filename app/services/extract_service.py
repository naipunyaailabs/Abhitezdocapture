import os
import io
import uuid
import base64
from typing import Dict, Any, Tuple
from app.utils.pdf_parser import pdf_parser
from app.services.llm_service import llm_service
import json

class ExtractService:
    async def extract_doc(self, buffer: bytes, file_name: str, file_type: str) -> str:
        ext = file_name.split(".")[-1].lower() if "." in file_name else ""
        print(f"[extract_doc] file_name: {file_name}, ext: {ext}, file_type: {file_type}")
        
        if (file_type and file_type == "application/pdf") or ext == "pdf":
            return await self.extract_pdf_with_ocr_fallback(buffer)
        
        if (file_type and file_type.startswith("image/")) or ext in ["jpg", "jpeg", "png"]:
            return await self.ocr_image(buffer)

        if ext in ["doc", "docx"] or (file_type and "wordprocessingml" in file_type):
            return await self.extract_word_content(buffer)

        if (file_type and file_type == "text/plain") or ext == "txt":
            try:
                return buffer.decode('utf-8')
            except:
                return buffer.decode('latin-1', errors='ignore')
            
        # Fallback for other files
        b64 = base64.b64encode(buffer[:4000]).decode('utf-8')
        sys_prompt = f"You are an intelligent document parser. Extract ALL text and structure from this {file_type} file."
        usr_prompt = f"Document base64 (truncated): {b64}"
        return await llm_service.unified_chat_completion(sys_prompt, usr_prompt)

    async def extract_word_content(self, buffer: bytes) -> str:
        try:
            from docx import Document
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(buffer)
                tmp_path = tmp.name
            
            doc = Document(tmp_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    table_text.append(" | ".join(cell.text.strip() for cell in row.cells))
            
            os.unlink(tmp_path)
            return "\n\n".join(paragraphs + table_text)
        except Exception as e:
            print(f"Word Extraction Error: {e}")
            return ""

    async def ocr_image(self, buffer: bytes) -> str:
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(io.BytesIO(buffer))
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            print(f"OCR Error: {e}")
            return ""

    async def extract_pdf_with_ocr_fallback(self, buffer: bytes) -> str:
        # Try direct text extraction first
        text = await pdf_parser.extract_text(buffer)
        
        # Heuristic: if very little text, try OCR
        if len(text.strip()) < 50:
            print("[ExtractService] PDF has little text, attempting OCR fallback")
            images = await pdf_parser.pdf_to_images(buffer)
            ocr_text_parts = []
            for img_bytes in images:
                page_text = await self.ocr_image(img_bytes)
                ocr_text_parts.append(page_text)
            text = "\n\n".join(ocr_text_parts)

            # If OCR also failed (e.g. Tesseract not installed), use vision LLM directly
            if len(text.strip()) < 50 and images:
                print("[ExtractService] OCR unavailable, using vision LLM fallback")
                vision_parts = []
                for img_bytes in images:
                    img_b64 = base64.b64encode(img_bytes).decode("utf-8")
                    page_text = await llm_service.unified_chat_completion(
                        "You are a document text extractor. Extract ALL visible text from this document image exactly as it appears. Return only the extracted text.",
                        "Extract all text from this document image.",
                        image_base64=img_b64,
                        image_mime_type="image/png"
                    )
                    vision_parts.append(page_text)
                text = "\n\n".join(vision_parts)

        return text.strip()

    async def structured_extraction(self, text: str) -> Dict[str, Any]:
        system_prompt = """You are an Expert Forensic Data Extractor specializing in Goods Receipt and Invoice processing.
Your Mission: Perform a zero-loss, exhaustive extraction of all data points from the provided document, specifically targeting fields for a Goods Receipt template.

CRITICAL RULES:
1. **NO PLACEHOLDERS**: NEVER use the word "Mandatory" or "Required" as a value. If a piece of information is not found in the text, use "N/A".
2. **DATA INTEGRITY**: Perform a zero-loss extraction. If it's on the page, it must be in the JSON.
3. **LINE ITEM PRECISION**: 'item_amount' refers to the TOTAL for that specific line (Qty * Rate). DO NOT put the Grand Total of the entire invoice here.
4. **Specific Details**: Rigorously search for Lot No, Shade No, Item Code, and PO Number.
5. **Format Precision**: Output ONLY valid JSON.
6. **Data Normalization**: Format dates as YYYY-MM-DD. Normalize numbers to strings without extra text (e.g., "163" instead of "163 boxes")."""

        schema_instructions = """
        Target JSON Schema:
        {
            "supplier_name": "Full name of the seller/supplier",
            "gst_no": "GSTIN of the supplier",
            "invoice_no": "Invoice number",
            "invoice_date": "Date of the invoice (YYYY-MM-DD)",
            "challan_no": "Challan or Delivery Note number if present",
            "challan_date": "Date of the challan if present",
            "gate_entry_no": "Gate entry number if present",
            "gate_entry_date": "Date of gate entry if present",
            "po_number": "Purchase Order (PO) number",
            "line_items": [
                {
                    "item_code": "Code or SKU for the item if present",
                    "item_description": "Full description of the goods",
                    "lot_no": "Lot or Batch number",
                    "shade_no": "Shade or Color code if present",
                    "quantity": "Numeric quantity value",
                    "rate": "Unit price/rate",
                    "gst": "GST percentage or amount for this item",
                    "item_amount": "Total amount for this specific line item"
                }
            ],
            "other_details": {
                "additional descriptive key-value pairs for any data not covered above"
            }
        }
        """
        user_prompt = f"Perform an exhaustive, detailed extraction of ALL data points from this text:\n\n{text[:20000]}\n\n{schema_instructions}"
        
        raw_response = await llm_service.unified_chat_completion(system_prompt, user_prompt)
        
        # Try to parse JSON more robustly
        try:
            import re
            # Try to find content between ```json and ```
            json_match = re.search(r'```json\s*(.*?)\s*```', raw_response, re.DOTALL)
            if json_match:
                clean_json = json_match.group(1).strip()
            else:
                # Try to find content between ``` and ```
                json_match = re.search(r'```\s*(.*?)\s*```', raw_response, re.DOTALL)
                if json_match:
                    clean_json = json_match.group(1).strip()
                else:
                    # Last resort: find the first { and last }
                    json_match = re.search(r'({.*})', raw_response, re.DOTALL)
                    if json_match:
                        clean_json = json_match.group(1).strip()
                    else:
                        clean_json = raw_response.strip()
            
            print(f"[DEBUG] Cleaned JSON for parsing: {clean_json[:500]}...")
            return json.loads(clean_json)
        except Exception as e:
            print(f"Structured extraction parse error: {e}")
            print(f"[DEBUG] Failed to parse JSON. Raw response: {raw_response[:500]}...")
            return {"error": "JSON parse error", "raw_text": raw_response}

    async def extract_and_parse(self, buffer: bytes, file_name: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        text = await self.extract_doc(buffer, file_name, file_type)
        structured = await self.structured_extraction(text)
        return text, structured

extract_service = ExtractService()
